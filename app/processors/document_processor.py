"""
Document Processor - Handles various document formats
"""

import io
import re
import hashlib
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional, Union, BinaryIO
from datetime import datetime
from enum import Enum


class DocumentType(Enum):
    """Supported document types."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MD = "markdown"
    HTML = "html"
    CSV = "csv"
    XLSX = "xlsx"
    JSON = "json"
    UNKNOWN = "unknown"


@dataclass
class DocumentMetadata:
    """Document metadata."""
    filename: str
    doc_type: DocumentType
    size_bytes: int
    page_count: int = 1
    word_count: int = 0
    char_count: int = 0
    created_at: datetime = field(default_factory=datetime.now)
    author: Optional[str] = None
    title: Optional[str] = None
    language: str = "en"
    checksum: str = ""
    custom_metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class DocumentChunk:
    """A chunk of document content."""
    content: str
    chunk_id: str
    document_id: str
    page_number: int = 1
    chunk_index: int = 0
    start_char: int = 0
    end_char: int = 0
    metadata: Dict[str, Any] = field(default_factory=dict)
    embedding: Optional[List[float]] = None

    def to_dict(self) -> Dict[str, Any]:
        return {
            "content": self.content,
            "chunk_id": self.chunk_id,
            "document_id": self.document_id,
            "page_number": self.page_number,
            "chunk_index": self.chunk_index,
            "metadata": self.metadata
        }


@dataclass
class ProcessedDocument:
    """Fully processed document."""
    document_id: str
    metadata: DocumentMetadata
    raw_content: str
    chunks: List[DocumentChunk]
    tables: List[Dict[str, Any]] = field(default_factory=list)
    images: List[Dict[str, Any]] = field(default_factory=list)
    sections: List[Dict[str, Any]] = field(default_factory=list)


class BaseDocumentProcessor(ABC):
    """Abstract base class for document processors."""

    @abstractmethod
    def process(self, content: Union[str, bytes, BinaryIO], filename: str) -> ProcessedDocument:
        """Process document content."""
        pass

    @abstractmethod
    def extract_text(self, content: Union[str, bytes, BinaryIO]) -> str:
        """Extract plain text from document."""
        pass

    def generate_document_id(self, content: bytes, filename: str) -> str:
        """Generate unique document ID."""
        hash_input = f"{filename}{len(content)}{datetime.now().isoformat()}"
        return hashlib.sha256(hash_input.encode()).hexdigest()[:16]


class TextDocumentProcessor(BaseDocumentProcessor):
    """Process plain text documents."""

    def process(self, content: Union[str, bytes, BinaryIO], filename: str) -> ProcessedDocument:
        text = self.extract_text(content)
        content_bytes = text.encode() if isinstance(content, str) else content

        doc_id = self.generate_document_id(
            content_bytes if isinstance(content_bytes, bytes) else b'',
            filename
        )

        metadata = DocumentMetadata(
            filename=filename,
            doc_type=DocumentType.TXT,
            size_bytes=len(content_bytes) if isinstance(content_bytes, bytes) else len(text),
            word_count=len(text.split()),
            char_count=len(text),
            checksum=hashlib.md5(text.encode()).hexdigest()
        )

        # Create chunks
        chunks = self._create_chunks(text, doc_id)

        return ProcessedDocument(
            document_id=doc_id,
            metadata=metadata,
            raw_content=text,
            chunks=chunks
        )

    def extract_text(self, content: Union[str, bytes, BinaryIO]) -> str:
        if isinstance(content, str):
            return content
        elif isinstance(content, bytes):
            return content.decode('utf-8', errors='ignore')
        else:
            return content.read().decode('utf-8', errors='ignore')

    def _create_chunks(
        self,
        text: str,
        doc_id: str,
        chunk_size: int = 1000,
        overlap: int = 200
    ) -> List[DocumentChunk]:
        """Split text into overlapping chunks."""
        chunks = []

        # Split by paragraphs first
        paragraphs = text.split('\n\n')

        current_chunk = ""
        current_start = 0
        chunk_index = 0

        for para in paragraphs:
            if len(current_chunk) + len(para) < chunk_size:
                current_chunk += para + "\n\n"
            else:
                if current_chunk.strip():
                    chunks.append(DocumentChunk(
                        content=current_chunk.strip(),
                        chunk_id=f"{doc_id}_chunk_{chunk_index}",
                        document_id=doc_id,
                        chunk_index=chunk_index,
                        start_char=current_start,
                        end_char=current_start + len(current_chunk)
                    ))
                    chunk_index += 1

                # Start new chunk with overlap
                overlap_text = current_chunk[-overlap:] if len(current_chunk) > overlap else ""
                current_start = current_start + len(current_chunk) - len(overlap_text)
                current_chunk = overlap_text + para + "\n\n"

        # Add final chunk
        if current_chunk.strip():
            chunks.append(DocumentChunk(
                content=current_chunk.strip(),
                chunk_id=f"{doc_id}_chunk_{chunk_index}",
                document_id=doc_id,
                chunk_index=chunk_index,
                start_char=current_start,
                end_char=current_start + len(current_chunk)
            ))

        return chunks


class MarkdownProcessor(TextDocumentProcessor):
    """Process Markdown documents."""

    def process(self, content: Union[str, bytes, BinaryIO], filename: str) -> ProcessedDocument:
        doc = super().process(content, filename)
        doc.metadata.doc_type = DocumentType.MD

        # Extract sections
        doc.sections = self._extract_sections(doc.raw_content)

        return doc

    def _extract_sections(self, text: str) -> List[Dict[str, Any]]:
        """Extract markdown sections by headers."""
        sections = []
        current_section = {"title": "Introduction", "level": 0, "content": ""}

        for line in text.split('\n'):
            header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if header_match:
                if current_section["content"].strip():
                    sections.append(current_section)

                level = len(header_match.group(1))
                title = header_match.group(2)
                current_section = {"title": title, "level": level, "content": ""}
            else:
                current_section["content"] += line + "\n"

        if current_section["content"].strip():
            sections.append(current_section)

        return sections


class PDFProcessor(BaseDocumentProcessor):
    """Process PDF documents."""

    def process(self, content: Union[str, bytes, BinaryIO], filename: str) -> ProcessedDocument:
        text = self.extract_text(content)

        if isinstance(content, bytes):
            content_bytes = content
        elif hasattr(content, 'read'):
            content_bytes = content.read()
            content.seek(0)
        else:
            content_bytes = content.encode()

        doc_id = self.generate_document_id(content_bytes, filename)

        # Count pages (simplified - in production use PyPDF2)
        page_count = text.count('\f') + 1 if '\f' in text else 1

        metadata = DocumentMetadata(
            filename=filename,
            doc_type=DocumentType.PDF,
            size_bytes=len(content_bytes),
            page_count=page_count,
            word_count=len(text.split()),
            char_count=len(text),
            checksum=hashlib.md5(content_bytes).hexdigest()
        )

        # Create chunks with page awareness
        chunks = self._create_chunks_with_pages(text, doc_id)

        return ProcessedDocument(
            document_id=doc_id,
            metadata=metadata,
            raw_content=text,
            chunks=chunks
        )

    def extract_text(self, content: Union[str, bytes, BinaryIO]) -> str:
        """Extract text from PDF using pypdf."""
        try:
            from pypdf import PdfReader

            if isinstance(content, bytes):
                pdf_file = io.BytesIO(content)
            elif isinstance(content, str):
                # Assume it's a file path
                pdf_file = open(content, 'rb')
            else:
                pdf_file = content

            reader = PdfReader(pdf_file)
            text_parts = []

            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                text_parts.append(f"[Page {page_num + 1}]\n{page_text}")

            return "\n\n".join(text_parts)

        except ImportError:
            return "[PDF processing requires pypdf library]"
        except Exception as e:
            return f"[Error extracting PDF text: {str(e)}]"

    def _create_chunks_with_pages(
        self,
        text: str,
        doc_id: str,
        chunk_size: int = 1000
    ) -> List[DocumentChunk]:
        """Create chunks while tracking page numbers."""
        chunks = []

        # Split by page markers
        page_pattern = r'\[Page (\d+)\]'
        pages = re.split(page_pattern, text)

        chunk_index = 0
        for i in range(1, len(pages), 2):
            page_num = int(pages[i])
            page_content = pages[i + 1] if i + 1 < len(pages) else ""

            # Split page content into chunks
            for j in range(0, len(page_content), chunk_size):
                chunk_text = page_content[j:j + chunk_size].strip()
                if chunk_text:
                    chunks.append(DocumentChunk(
                        content=chunk_text,
                        chunk_id=f"{doc_id}_chunk_{chunk_index}",
                        document_id=doc_id,
                        page_number=page_num,
                        chunk_index=chunk_index,
                        metadata={"page": page_num}
                    ))
                    chunk_index += 1

        return chunks


class DocxProcessor(BaseDocumentProcessor):
    """Process Word documents."""

    def process(self, content: Union[str, bytes, BinaryIO], filename: str) -> ProcessedDocument:
        text, tables = self._extract_content(content)

        if isinstance(content, bytes):
            content_bytes = content
        elif hasattr(content, 'read'):
            content_bytes = content.read()
            content.seek(0)
        else:
            content_bytes = content.encode()

        doc_id = self.generate_document_id(content_bytes, filename)

        metadata = DocumentMetadata(
            filename=filename,
            doc_type=DocumentType.DOCX,
            size_bytes=len(content_bytes),
            word_count=len(text.split()),
            char_count=len(text),
            checksum=hashlib.md5(content_bytes).hexdigest()
        )

        # Create chunks
        processor = TextDocumentProcessor()
        chunks = processor._create_chunks(text, doc_id)

        return ProcessedDocument(
            document_id=doc_id,
            metadata=metadata,
            raw_content=text,
            chunks=chunks,
            tables=tables
        )

    def extract_text(self, content: Union[str, bytes, BinaryIO]) -> str:
        text, _ = self._extract_content(content)
        return text

    def _extract_content(self, content: Union[str, bytes, BinaryIO]) -> tuple:
        """Extract text and tables from DOCX."""
        try:
            from docx import Document

            if isinstance(content, bytes):
                doc_file = io.BytesIO(content)
            elif isinstance(content, str):
                doc_file = content
            else:
                doc_file = content

            doc = Document(doc_file)

            # Extract paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    # Check for heading style
                    if para.style.name.startswith('Heading'):
                        level = para.style.name[-1] if para.style.name[-1].isdigit() else '1'
                        text_parts.append(f"\n{'#' * int(level)} {para.text}\n")
                    else:
                        text_parts.append(para.text)

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append({
                    "rows": len(table.rows),
                    "cols": len(table.columns),
                    "data": table_data
                })

            return "\n\n".join(text_parts), tables

        except ImportError:
            return "[DOCX processing requires python-docx library]", []
        except Exception as e:
            return f"[Error extracting DOCX: {str(e)}]", []


class CSVProcessor(BaseDocumentProcessor):
    """Process CSV documents."""

    def process(self, content: Union[str, bytes, BinaryIO], filename: str) -> ProcessedDocument:
        text = self.extract_text(content)
        tables = self._parse_csv(content)

        if isinstance(content, bytes):
            content_bytes = content
        elif isinstance(content, str):
            content_bytes = content.encode()
        else:
            content_bytes = content.read()
            content.seek(0)

        doc_id = self.generate_document_id(content_bytes, filename)

        metadata = DocumentMetadata(
            filename=filename,
            doc_type=DocumentType.CSV,
            size_bytes=len(content_bytes),
            word_count=len(text.split()),
            char_count=len(text),
            checksum=hashlib.md5(content_bytes).hexdigest()
        )

        # Create chunks from rows
        chunks = self._create_row_chunks(tables, doc_id)

        return ProcessedDocument(
            document_id=doc_id,
            metadata=metadata,
            raw_content=text,
            chunks=chunks,
            tables=tables
        )

    def extract_text(self, content: Union[str, bytes, BinaryIO]) -> str:
        if isinstance(content, bytes):
            return content.decode('utf-8', errors='ignore')
        elif hasattr(content, 'read'):
            return content.read().decode('utf-8', errors='ignore')
        return content

    def _parse_csv(self, content: Union[str, bytes, BinaryIO]) -> List[Dict[str, Any]]:
        """Parse CSV into structured table."""
        import csv

        text = self.extract_text(content)
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)

        if not rows:
            return []

        headers = rows[0] if rows else []
        data = rows[1:] if len(rows) > 1 else []

        return [{
            "headers": headers,
            "rows": len(data),
            "cols": len(headers),
            "data": data
        }]

    def _create_row_chunks(
        self,
        tables: List[Dict[str, Any]],
        doc_id: str,
        rows_per_chunk: int = 50
    ) -> List[DocumentChunk]:
        """Create chunks from table rows."""
        chunks = []

        for table in tables:
            headers = table.get("headers", [])
            data = table.get("data", [])

            for i in range(0, len(data), rows_per_chunk):
                chunk_rows = data[i:i + rows_per_chunk]

                # Format as text
                content_parts = [" | ".join(headers)]
                content_parts.append("-" * 50)
                for row in chunk_rows:
                    content_parts.append(" | ".join(str(cell) for cell in row))

                chunks.append(DocumentChunk(
                    content="\n".join(content_parts),
                    chunk_id=f"{doc_id}_chunk_{len(chunks)}",
                    document_id=doc_id,
                    chunk_index=len(chunks),
                    metadata={
                        "row_start": i,
                        "row_end": i + len(chunk_rows),
                        "headers": headers
                    }
                ))

        return chunks


class DocumentProcessorFactory:
    """Factory for creating document processors."""

    _processors = {
        DocumentType.TXT: TextDocumentProcessor,
        DocumentType.MD: MarkdownProcessor,
        DocumentType.PDF: PDFProcessor,
        DocumentType.DOCX: DocxProcessor,
        DocumentType.CSV: CSVProcessor,
    }

    @classmethod
    def get_processor(cls, doc_type: DocumentType) -> BaseDocumentProcessor:
        """Get processor for document type."""
        processor_class = cls._processors.get(doc_type, TextDocumentProcessor)
        return processor_class()

    @classmethod
    def detect_type(cls, filename: str) -> DocumentType:
        """Detect document type from filename."""
        ext = filename.lower().split('.')[-1]
        type_map = {
            'pdf': DocumentType.PDF,
            'docx': DocumentType.DOCX,
            'doc': DocumentType.DOCX,
            'txt': DocumentType.TXT,
            'md': DocumentType.MD,
            'markdown': DocumentType.MD,
            'csv': DocumentType.CSV,
            'xlsx': DocumentType.XLSX,
            'html': DocumentType.HTML,
            'htm': DocumentType.HTML,
            'json': DocumentType.JSON,
        }
        return type_map.get(ext, DocumentType.UNKNOWN)

    @classmethod
    def process_document(
        cls,
        content: Union[str, bytes, BinaryIO],
        filename: str
    ) -> ProcessedDocument:
        """Process document with automatic type detection."""
        doc_type = cls.detect_type(filename)
        processor = cls.get_processor(doc_type)
        return processor.process(content, filename)
